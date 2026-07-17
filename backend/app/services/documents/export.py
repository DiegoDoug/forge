from __future__ import annotations

import io
import re
from datetime import datetime
from xml.sax.saxutils import escape as xml_escape

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from markdownify import markdownify
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .html_blocks import parse_blocks

SUPPORTED_FORMATS = ("txt", "md", "xml", "doc", "docx", "pdf")

_UNSAFE_FILENAME = re.compile(r"[^A-Za-z0-9._ -]+")


def safe_filename(title: str, ext: str) -> str:
    name = (title or "Untitled").strip()
    name = _UNSAFE_FILENAME.sub("_", name).strip(" .") or "Untitled"
    return f"{name[:120]}.{ext}"


# --- .txt -------------------------------------------------------------------


def to_txt(html: str) -> str:
    blocks = parse_blocks(html)
    if not blocks:
        return "\n"
    lines = []
    for block in blocks:
        text = "".join(r["text"] for r in block["runs"]).strip("\n")
        if block["type"] == "list_item":
            bullet = f"{block.get('index') or 1}. " if block.get("ordered") else "- "
            text = bullet + text
        lines.append(text)
    return "\n\n".join(lines).rstrip() + "\n"


# --- .md ----------------------------------------------------------------------


_STRIKE_TAG_RE = re.compile(r"<(/?)strike>", re.IGNORECASE)


def to_markdown(html: str) -> str:
    # markdownify only recognizes <s>/<del> for ~~strikethrough~~ — the
    # editor's strikeThrough command emits the legacy <strike> tag.
    normalized = _STRIKE_TAG_RE.sub(r"<\1s>", html or "")
    md = markdownify(normalized, heading_style="ATX", bullets="-").strip()
    return (md or "") + "\n"


# --- .xml ---------------------------------------------------------------------


def _cdata(text: str) -> str:
    return (text or "").replace("]]>", "]]]]><![CDATA[>")


def to_xml(title: str, html: str, created_at: datetime, updated_at: datetime) -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        "<document>\n"
        f"  <title>{xml_escape(title or 'Untitled')}</title>\n"
        f"  <created_at>{created_at.isoformat()}</created_at>\n"
        f"  <updated_at>{updated_at.isoformat()}</updated_at>\n"
        f"  <content format=\"html\"><![CDATA[{_cdata(html or '')}]]></content>\n"
        f"  <content format=\"text\"><![CDATA[{_cdata(to_txt(html))}]]></content>\n"
        "</document>\n"
    )


# --- .doc (RTF) -----------------------------------------------------------
# Genuine RTF is what Word/LibreOffice/WordPad/Pages actually expect from a
# ".doc" file produced outside real Word — this is the same trick classic
# "export to Word" features have long used, and unlike raw OLE .doc it can be
# generated as plain text without a binary-format library.

_DEFAULT_FONT = "Calibri"


def _rtf_escape(text: str) -> str:
    out = []
    for ch in text:
        code = ord(ch)
        if ch == "\\":
            out.append("\\\\")
        elif ch == "{":
            out.append("\\{")
        elif ch == "}":
            out.append("\\}")
        elif ch == "\n":
            out.append("\\line ")
        elif code < 128:
            out.append(ch)
        elif code <= 0xFFFF:
            out.append(f"\\u{code}?")
        else:
            out.append("?")
    return "".join(out)


_RTF_ALIGN = {"left": "\\ql", "center": "\\qc", "right": "\\qr", "justify": "\\qj"}
_RTF_HEADING_HALF_PT = {1: 64, 2: 56, 3: 48, 4: 44, 5: 40, 6: 36}


def to_rtf(title: str, html: str) -> bytes:
    blocks = parse_blocks(html)
    fonts = [_DEFAULT_FONT]

    def font_index(name: str | None) -> int:
        name = name or _DEFAULT_FONT
        if name not in fonts:
            fonts.append(name)
        return fonts.index(name)

    paragraphs = []

    if title:
        paragraphs.append("{\\pard\\qc\\b\\fs36 " + _rtf_escape(title) + "\\b0\\par}")

    for block in blocks:
        align = _RTF_ALIGN.get(block["align"], "\\ql")
        control = ["\\pard", align]
        is_heading = block["type"] == "heading"
        run_parts = []
        if block["type"] == "list_item":
            bullet = f"{block.get('index') or 1}. " if block.get("ordered") else "\\bullet  "
            control.append("\\li360")
            run_parts.append(bullet if bullet.startswith("\\") else _rtf_escape(bullet))

        # Headings get a bold+larger-size floor; per-run formatting from the
        # HTML (explicit size/bold on a nested tag) still wins when present —
        # applying it *inside* each run (rather than once for the whole
        # paragraph) means it can't be clobbered by a run's own \fs control.
        default_size = (_RTF_HEADING_HALF_PT.get(block.get("level") or 1, 32) / 2) if is_heading else 12

        for run in block["runs"]:
            bold = bool(run.get("bold")) or is_heading
            codes = []
            if bold:
                codes.append("\\b")
            if run.get("italic"):
                codes.append("\\i")
            if run.get("underline"):
                codes.append("\\ul")
            if run.get("strike"):
                codes.append("\\strike")
            fidx = font_index(run.get("font"))
            codes.append(f"\\f{fidx}")
            size_pt = run.get("size") or default_size
            codes.append(f"\\fs{int(round(size_pt * 2))}")
            reset = []
            if bold:
                reset.append("\\b0")
            if run.get("italic"):
                reset.append("\\i0")
            if run.get("underline"):
                reset.append("\\ulnone")
            if run.get("strike"):
                reset.append("\\strike0")
            run_parts.append("{" + "".join(codes) + " " + _rtf_escape(run["text"]) + "".join(reset) + "}")
        paragraphs.append("{" + "".join(control) + " " + "".join(run_parts) + "\\par}")

    font_table = "".join(f"{{\\f{i} {name};}}" for i, name in enumerate(fonts))
    body = "\n".join(paragraphs) if paragraphs else "{\\pard\\par}"
    rtf = (
        "{\\rtf1\\ansi\\ansicpg1252\\deff0\\deflang1033\n"
        f"{{\\fonttbl{font_table}}}\n"
        "\\fs24\n"
        f"{body}\n"
        "}"
    )
    return rtf.encode("cp1252", errors="replace")


# --- .docx ----------------------------------------------------------------

_DOCX_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def to_docx(title: str, html: str) -> bytes:
    blocks = parse_blocks(html)
    doc = DocxDocument()

    if title:
        doc.add_heading(title, level=0)

    for block in blocks:
        if block["type"] == "heading":
            paragraph = doc.add_heading("", level=min(max(block.get("level") or 1, 1), 9))
        elif block["type"] == "list_item":
            style = "List Number" if block.get("ordered") else "List Bullet"
            paragraph = doc.add_paragraph(style=style)
        else:
            paragraph = doc.add_paragraph()
        paragraph.alignment = _DOCX_ALIGN.get(block["align"], WD_ALIGN_PARAGRAPH.LEFT)

        for run_data in block["runs"]:
            text = run_data["text"]
            if text == "":
                continue
            run = paragraph.add_run(text)
            run.bold = run_data.get("bold") or None
            run.italic = run_data.get("italic") or None
            run.underline = run_data.get("underline") or None
            if run_data.get("strike"):
                run.font.strike = True
            if run_data.get("font"):
                run.font.name = run_data["font"]
            if run_data.get("size"):
                run.font.size = Pt(run_data["size"])

    if not blocks and not title:
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- .pdf -------------------------------------------------------------------

_PDF_FAMILY_MAP = {
    "arial": "Helvetica",
    "helvetica": "Helvetica",
    "verdana": "Helvetica",
    "sans-serif": "Helvetica",
    "tahoma": "Helvetica",
    "times new roman": "Times-Roman",
    "times": "Times-Roman",
    "georgia": "Times-Roman",
    "serif": "Times-Roman",
    "courier new": "Courier",
    "courier": "Courier",
    "monospace": "Courier",
}

_PDF_ALIGN = {"left": TA_LEFT, "center": TA_CENTER, "right": TA_RIGHT, "justify": TA_JUSTIFY}


def _pdf_family(name: str | None) -> str:
    return _PDF_FAMILY_MAP.get((name or "").strip().lower(), "Helvetica")


def _run_markup(run: dict) -> str:
    text = xml_escape(run["text"]).replace("\n", "<br/>")
    if not text:
        return ""
    if run.get("bold"):
        text = f"<b>{text}</b>"
    if run.get("italic"):
        text = f"<i>{text}</i>"
    if run.get("underline"):
        text = f"<u>{text}</u>"
    if run.get("strike"):
        text = f"<strike>{text}</strike>"
    attrs = f' name="{_pdf_family(run.get("font"))}"'
    if run.get("size"):
        attrs += f' size="{run["size"]}"'
    return f"<font{attrs}>{text}</font>"


_PDF_HEADING_SIZE = {1: 22, 2: 19, 3: 16, 4: 14, 5: 12.5, 6: 11.5}


def to_pdf(title: str, html: str) -> bytes:
    blocks = parse_blocks(html)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, topMargin=56, bottomMargin=56, leftMargin=56, rightMargin=56)
    story = []

    if title:
        title_style = ParagraphStyle("DocTitle", fontName="Helvetica-Bold", fontSize=22, leading=26, spaceAfter=16)
        story.append(Paragraph(xml_escape(title), title_style))

    base_style = ParagraphStyle("Body", fontName="Helvetica", fontSize=11, leading=15.5, spaceAfter=8)

    for block in blocks:
        markup = "".join(_run_markup(r) for r in block["runs"]) or "&nbsp;"
        align = _PDF_ALIGN.get(block["align"], TA_LEFT)
        if block["type"] == "heading":
            size = _PDF_HEADING_SIZE.get(block.get("level") or 1, 14)
            style = ParagraphStyle(
                f"H{block.get('level') or 1}-{id(block)}",
                fontName="Helvetica-Bold",
                fontSize=size,
                leading=size * 1.25,
                spaceBefore=10,
                spaceAfter=6,
                alignment=align,
            )
            story.append(Paragraph(markup, style))
        elif block["type"] == "list_item":
            style = ParagraphStyle(
                f"List-{id(block)}", parent=base_style, alignment=align, leftIndent=20, bulletIndent=6
            )
            bullet = f"{block.get('index') or 1}." if block.get("ordered") else "•"
            story.append(Paragraph(markup, style, bulletText=bullet))
        else:
            style = ParagraphStyle(f"P-{id(block)}", parent=base_style, alignment=align)
            story.append(Paragraph(markup, style))

    if not story:
        story.append(Spacer(1, 1))

    doc.build(story)
    return buf.getvalue()


# --- dispatch -----------------------------------------------------------------

_MEDIA_TYPES = {
    "txt": "text/plain; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "xml": "application/xml; charset=utf-8",
    "doc": "application/msword",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def export_document(*, title: str, content: str, fmt: str, created_at: datetime, updated_at: datetime) -> tuple[bytes, str, str]:
    """Returns (payload, media_type, filename) for the requested export format."""

    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported export format: {fmt}")

    if fmt == "txt":
        payload = to_txt(content).encode("utf-8")
    elif fmt == "md":
        payload = to_markdown(content).encode("utf-8")
    elif fmt == "xml":
        payload = to_xml(title, content, created_at, updated_at).encode("utf-8")
    elif fmt == "doc":
        payload = to_rtf(title, content)
    elif fmt == "docx":
        payload = to_docx(title, content)
    elif fmt == "pdf":
        payload = to_pdf(title, content)
    else:  # pragma: no cover — guarded above
        raise ValueError(f"Unsupported export format: {fmt}")

    return payload, _MEDIA_TYPES[fmt], safe_filename(title, fmt)
