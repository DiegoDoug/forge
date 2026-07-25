"""Baseline characterization tests for the Ingest pipeline.

Written as part of Universal Converter (Phase 04) Milestone 1, before any
production code in this phase changes. `services/ingest/` has shipped with
zero test coverage until now — these tests pin its current, real behavior
(observed by running the actual code, not guessed) so that Milestone 3's
`MarkItDownProvider` wrapper has something concrete to be checked against.
Per `03_BACKEND.md` §1, that wrapper must delegate to this exact pipeline
without changing its behavior; these tests are the mechanism that makes that
claim falsifiable rather than assumed.
"""

from __future__ import annotations

import zipfile

import pytest

from app.services.ingest import formats
from app.services.ingest.jobs import FileTask, Job, JobStore
from app.services.ingest.postprocess import clean_markdown

# --- postprocess.clean_markdown --------------------------------------------


def test_clean_markdown_strips_invisible_characters():
    assert clean_markdown("Hello​World\n") == "HelloWorld\n"


def test_clean_markdown_normalizes_smart_quotes_and_dashes():
    assert clean_markdown("‘Hi’ “there”—ok…\n") == "'Hi' \"there\" - ok...\n"


def test_clean_markdown_collapses_long_data_uri_images():
    long_data_uri = "![alt](data:image/png;base64," + "A" * 70 + ")\n"
    assert clean_markdown(long_data_uri) == "![alt](embedded-image)\n"


def test_clean_markdown_collapses_multiple_blank_lines_to_one():
    assert clean_markdown("Line1\n\n\n\nLine2\n") == "Line1\n\nLine2\n"


def test_clean_markdown_preserves_blank_lines_inside_code_fences():
    fenced = "```\ncode1\n\n\ncode2\n```\n"
    assert clean_markdown(fenced) == fenced


def test_clean_markdown_empty_input_returns_empty_string():
    assert clean_markdown("") == ""


# --- formats.KNOWN_EXTENSIONS -----------------------------------------------


def test_known_extensions_snapshot():
    # A deliberate snapshot assertion: any future change to the supported
    # extension list must be a visible, reviewed diff here, not a silent
    # drop or addition.
    assert formats.KNOWN_EXTENSIONS == {
        ".pdf", ".docx", ".doc", ".pptx", ".ppt", ".xlsx", ".xls", ".csv",
        ".html", ".htm", ".xhtml", ".epub", ".json", ".jsonl", ".xml", ".rss",
        ".atom", ".txt", ".md", ".markdown", ".rst",
        ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp",
        ".wav", ".mp3", ".m4a", ".mp4", ".flac", ".ogg",
        ".zip", ".msg", ".ipynb", ".yaml", ".yml", ".log", ".tsv",
    }


# --- jobs.JobStore -----------------------------------------------------------


@pytest.fixture()
def job_store(tmp_path, monkeypatch):
    """A `JobStore` instance backed by an isolated tmp_path, not the module's
    real global `store` singleton and not the shared FORGE_DATA_DIR — so
    these tests can't collide with each other or with any other test module
    that touches Ingest's real, process-wide job store."""

    class _FakeSettings:
        ingest_jobs_dir = tmp_path / "ingest-jobs"
        ingest_job_ttl_minutes = 120

    import app.services.ingest.jobs as jobs_module

    monkeypatch.setattr(jobs_module, "get_settings", lambda: _FakeSettings())
    return JobStore()


def test_job_store_create_makes_in_and_out_directories(job_store):
    job = job_store.create()
    assert (job.dir() / "in").is_dir()
    assert (job.dir() / "out").is_dir()
    assert job_store.get(job.id) is job


def test_job_status_aggregates_file_statuses():
    job = Job(id="j1", created_at=0)
    job.files = [FileTask(id="a", original_name="a.txt", upload_path=None, status="done")]
    assert job.status == "done"

    job.files.append(FileTask(id="b", original_name="b.txt", upload_path=None, status="processing"))
    assert job.status == "processing"

    job.files[-1].status = "error"
    assert job.status == "done"  # mixed done/error still reports "done" overall

    job.files[0].status = "error"
    assert job.status == "failed"  # every file errored


def test_job_progress_is_finished_over_total():
    job = Job(id="j2", created_at=0)
    job.files = [
        FileTask(id="a", original_name="a.txt", upload_path=None, status="done"),
        FileTask(id="b", original_name="b.txt", upload_path=None, status="processing"),
        FileTask(id="c", original_name="c.txt", upload_path=None, status="error"),
        FileTask(id="d", original_name="d.txt", upload_path=None, status="pending"),
    ]
    assert job.progress == 0.5  # 2 of 4 finished (done + error)


def test_job_progress_with_no_files_is_one():
    job = Job(id="j3", created_at=0)
    assert job.progress == 1.0


def test_cleanup_expired_removes_only_jobs_past_ttl(job_store):
    fresh = job_store.create()
    expired = job_store.create()
    expired.created_at = 0  # far in the past — guaranteed expired regardless of TTL

    job_store.cleanup_expired()

    assert job_store.get(fresh.id) is not None
    assert (fresh.dir()).exists()
    assert job_store.get(expired.id) is None
    assert not expired.dir().exists()


# --- converter._convert_one across format families --------------------------


def _make_job_and_task(tmp_path, name: str, content: bytes) -> tuple[Job, FileTask]:
    job = Job(id="convert-" + name.replace(".", "-"), created_at=0)
    in_dir = tmp_path / job.id / "in"
    out_dir = tmp_path / job.id / "out"
    in_dir.mkdir(parents=True)
    out_dir.mkdir(parents=True)
    upload_path = in_dir / name
    upload_path.write_bytes(content)
    task = FileTask(id="t1", original_name=name, upload_path=upload_path)
    job.files.append(task)
    return job, task


@pytest.fixture()
def isolated_job(tmp_path, monkeypatch):
    """Points `Job.dir()` at tmp_path for the conversion tests below, via the
    same settings-patch technique as `job_store` above."""

    class _FakeSettings:
        ingest_jobs_dir = tmp_path

    import app.services.ingest.jobs as jobs_module

    monkeypatch.setattr(jobs_module, "get_settings", lambda: _FakeSettings())
    return tmp_path


def test_convert_plain_text(isolated_job):
    from app.services.ingest import converter

    job, task = _make_job_and_task(isolated_job, "sample.txt", b"Hello Forge baseline test.\n")
    converter._convert_one(job, task)
    assert task.status == "done"
    assert task.output_path.read_text("utf-8") == "Hello Forge baseline test.\n"


def test_convert_html(isolated_job):
    from app.services.ingest import converter

    html = b"<html><body><h1>Title</h1><p>Body text here.</p></body></html>"
    job, task = _make_job_and_task(isolated_job, "sample.html", html)
    converter._convert_one(job, task)
    assert task.status == "done"
    assert task.output_path.read_text("utf-8") == "# Title\n\nBody text here.\n"


def test_convert_docx(isolated_job):
    import io

    from docx import Document as DocxDocument

    from app.services.ingest import converter

    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Sample Heading", level=1)
    doc.add_paragraph("Sample paragraph body text.")
    doc.save(buf)

    job, task = _make_job_and_task(isolated_job, "sample.docx", buf.getvalue())
    converter._convert_one(job, task)
    assert task.status == "done"
    assert task.output_path.read_text("utf-8") == "# Sample Heading\n\nSample paragraph body text.\n"


def test_convert_xlsx(isolated_job):
    import io

    from openpyxl import Workbook

    from app.services.ingest import converter

    buf = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Header"
    ws["A2"] = "Value123"
    wb.save(buf)

    job, task = _make_job_and_task(isolated_job, "sample.xlsx", buf.getvalue())
    converter._convert_one(job, task)
    assert task.status == "done"
    output = task.output_path.read_text("utf-8")
    assert "Header" in output
    assert "Value123" in output


def test_convert_pdf_without_vision(isolated_job, monkeypatch):
    import io

    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    from app.services.ingest import converter

    # Vision is disabled by default (FORGE_VISION_ENABLED unset) — confirm
    # that's really the config in effect for this test, since a real-text PDF
    # should convert via the plain (non-vision) path either way.
    from app.core.config import get_settings

    assert get_settings().vision_enabled is False

    buf = io.BytesIO()
    sdoc = SimpleDocTemplate(buf, pagesize=LETTER)
    style = ParagraphStyle("Body", fontName="Helvetica", fontSize=11)
    sdoc.build([Paragraph("This is a sample PDF paragraph with real extractable text content.", style)])

    job, task = _make_job_and_task(isolated_job, "sample.pdf", buf.getvalue())
    converter._convert_one(job, task)
    assert task.status == "done"
    assert task.used_vision is False
    assert "sample PDF paragraph with real extractable text content" in task.output_path.read_text("utf-8")


def test_convert_zip_recurses_into_contents(isolated_job):
    import io

    from app.services.ingest import converter

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("sample.txt", "Hello Forge baseline test.\n")

    job, task = _make_job_and_task(isolated_job, "sample.zip", buf.getvalue())
    converter._convert_one(job, task)
    assert task.status == "done"
    output = task.output_path.read_text("utf-8")
    assert "sample.txt" in output
    assert "Hello Forge baseline test." in output


def test_convert_image_without_vision_fails_with_no_text_content(isolated_job):
    """Pins a real, current gap: without vision assistance, an image file
    carries no extractable text, so conversion fails rather than producing
    empty or placeholder Markdown. This is today's actual behavior — this
    test exists to catch it changing silently, not to bless it as desired
    UX."""
    import io

    from PIL import Image

    from app.services.ingest import converter

    buf = io.BytesIO()
    Image.new("RGB", (64, 64), color=(120, 50, 200)).save(buf, "JPEG")

    job, task = _make_job_and_task(isolated_job, "sample.jpg", buf.getvalue())
    converter._convert_one(job, task)
    assert task.status == "error"
    assert task.error is not None
    assert task.output_path is None


def test_convert_deletes_upload_file_regardless_of_outcome(isolated_job):
    from app.services.ingest import converter

    job, task = _make_job_and_task(isolated_job, "sample.txt", b"content\n")
    upload_path = task.upload_path
    assert upload_path.exists()
    converter._convert_one(job, task)
    assert not upload_path.exists()


# --- vision.pdf_needs_vision (default-off pinning) ---------------------------


def test_pdf_needs_vision_is_false_by_default(tmp_path):
    """Confirms the common/default deployment path: with FORGE_VISION_ENABLED
    unset (the default, per `01_PRODUCT_PRINCIPLES.md` §1.2's self-hosted,
    cloud-optional principle), no PDF ever triggers the vision fallback,
    regardless of its content."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import SimpleDocTemplate

    from app.services.ingest import vision

    pdf_path = tmp_path / "blank.pdf"
    SimpleDocTemplate(str(pdf_path), pagesize=LETTER).build([])

    needs_vision, page_count = vision.pdf_needs_vision(pdf_path)
    assert needs_vision is False
